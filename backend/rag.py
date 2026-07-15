# backend/rag.py
import os
import re
import csv
import json
import traceback
from typing import List, Dict, Any
import chromadb
from google import genai
import contextvars
from db import is_postgres_active, get_conn, get_cursor, execute_sql

# Context variable to store active user context during RAG search
active_user_context = contextvars.ContextVar("active_user_context", default=None)
# Context variable to optionally restrict RAG search to a single document file
active_file_context = contextvars.ContextVar("active_file_context", default=None)

# Database path (relative to backend folder)
CHROMA_DB_PATH = os.path.join(os.path.dirname(__file__), "data", "chromadb")
COLLECTION_NAME = "knowledge_worker_rag"

_collection: Any = None

def get_chroma_collection() -> Any:
    """
    Get or initialize the ChromaDB persistent collection.
    """
    global _collection
    if _collection is not None:
        return _collection

    os.makedirs(CHROMA_DB_PATH, exist_ok=True)
    try:
        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
        _collection = client.get_or_create_collection(name=COLLECTION_NAME)
        return _collection
    except Exception as e:
        print(f"Error initializing ChromaDB: {e}")
        traceback.print_exc()
        return None

def init_rag():
    """
    Called at startup to initialize ChromaDB.
    """
    col = get_chroma_collection()
    if col is not None:
        print(f"RAG Engine Initialized. ChromaDB collection '{COLLECTION_NAME}' loaded. Total items: {col.count()}")
    else:
        print("RAG Engine Initialization Failed.")

def get_gemini_client():
    """
    Initialize Gemini API client.
    """
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key or api_key.startswith("your_") or len(api_key.strip()) < 10:
        return None
    try:
        return genai.Client(api_key=api_key.strip())
    except Exception as e:
        print(f"Error initializing Gemini client for RAG: {e}")
        return None

def get_gemini_embeddings_batch(texts: List[str]) -> List[List[float]]:
    """
    Generate embeddings for a batch of texts using Gemini API.
    Tries gemini-embedding-2 first, then text-embedding-004 as fallback.
    Provides automatic fallback to individual queries if batch results are incomplete.
    """
    if not texts:
        return []

    client = get_gemini_client()
    if not client:
        raise ValueError("Gemini API key is not configured or is invalid. Cannot generate embeddings.")

    # Embedding models to try in order
    EMBED_MODELS = ["gemini-embedding-2", "text-embedding-004"]
    all_embeddings = []

    for embed_model in EMBED_MODELS:
        all_embeddings = []
        try:
            # 1. Try batch embedding first (efficient, 1 API call)
            try:
                texts_any: Any = texts
                response = client.models.embed_content(
                    model=embed_model,
                    contents=texts_any
                )
                response_any: Any = response
                if hasattr(response_any, 'embeddings') and response_any.embeddings:
                    for emb in response_any.embeddings:
                        all_embeddings.append(emb.values)
            except Exception as e:
                print(f"[{embed_model}] Batch embedding failed: {e}. Trying individual...")

            # 2. If batching returned incorrect length or failed, fallback to individual
            if len(all_embeddings) != len(texts):
                all_embeddings = []
                print(f"[{embed_model}] Embedding count mismatch. Embedding items individually...")
                for text in texts:
                    try:
                        response = client.models.embed_content(
                            model=embed_model,
                            contents=text
                        )
                        response_any2: Any = response
                        if hasattr(response_any2, 'embedding') and response_any2.embedding is not None:
                            emb2: Any = response_any2.embedding
                            if hasattr(emb2, 'values') and emb2.values is not None:
                                all_embeddings.append(list(emb2.values))
                        elif hasattr(response_any2, 'embeddings') and response_any2.embeddings:
                            all_embeddings.append(list(response_any2.embeddings[0].values))
                        else:
                            raise ValueError("Could not extract embedding values.")
                    except Exception as e2:
                        print(f"[{embed_model}] Individual embedding failed: {e2}")
                        raise e2

            if len(all_embeddings) == len(texts):
                print(f"[RAG] Successfully generated {len(all_embeddings)} embeddings using {embed_model}")
                return all_embeddings

        except Exception as model_err:
            print(f"[RAG] Embedding model '{embed_model}' failed: {model_err}. Trying next model...")
            continue

    # All models failed
    if len(all_embeddings) == len(texts):
        return all_embeddings
    raise ValueError(f"All embedding models failed. Cannot generate embeddings for {len(texts)} texts.")




# ── Advanced Chunking ──────────────────────────────────────────────────────────

def _split_sentences(text: str) -> List[str]:
    """
    Split text into sentences using punctuation boundaries and paragraph breaks.
    Handles common abbreviations to avoid false splits (e.g. 'Mr.', 'U.S.').
    """
    # Protect common abbreviations
    abbrev_pattern = re.compile(
        r'\b(Mr|Mrs|Ms|Dr|Prof|Sr|Jr|vs|etc|i\.e|e\.g|U\.S|U\.K|Fig|No)\.',
        re.IGNORECASE
    )
    protected = abbrev_pattern.sub(lambda m: m.group(0).replace('.', '<PERIOD>'), text)

    # Split on sentence-ending punctuation followed by whitespace + capital letter
    raw_sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z"])', protected)

    # Restore protected periods
    sentences = [s.replace('<PERIOD>', '.').strip() for s in raw_sentences if s.strip()]
    return sentences


def chunk_text_smart(
    text: str,
    chunk_size: int = 800,
    overlap_sentences: int = 1,
    chunk_type: str = "text",
    page_num: int = 0
) -> List[Dict[str, Any]]:
    """
    Sentence-aware, paragraph-respecting text chunker.

    Improvements over naive character splitter:
    - Splits at sentence boundaries (never mid-sentence)
    - Respects paragraph breaks (\\n\\n) as hard split points
    - Overlaps by `overlap_sentences` sentences across chunk boundaries for context continuity
    - Returns enriched chunk dicts: {text, chunk_type, page_num}

    Args:
        text: Raw text to chunk.
        chunk_size: Target maximum character length per chunk.
        overlap_sentences: Number of sentences to repeat at the start of the next chunk.
        chunk_type: Label for this chunk ('text', 'table', 'header').
        page_num: PDF page number (0 for non-PDF sources).

    Returns:
        List of chunk dicts with keys: text, chunk_type, page_num
    """
    if not text or not text.strip():
        return []

    chunks: List[Dict[str, Any]] = []

    # Split into paragraphs first (double newline is a hard boundary)
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

    for paragraph in paragraphs:
        sentences = _split_sentences(paragraph)
        if not sentences:
            continue

        current_sentences: List[str] = []
        current_len = 0

        for sentence in sentences:
            sentence_len = len(sentence)

            # If adding this sentence exceeds chunk_size and we have content, flush
            if current_len + sentence_len > chunk_size and current_sentences:
                chunk_text_str = ' '.join(current_sentences).strip()
                if chunk_text_str:
                    chunks.append({
                        "text": chunk_text_str,
                        "chunk_type": chunk_type,
                        "page_num": page_num
                    })
                # Keep last `overlap_sentences` sentences as context for next chunk
                current_sentences = current_sentences[-overlap_sentences:] if overlap_sentences > 0 else []
                current_len = sum(len(s) for s in current_sentences)

            current_sentences.append(sentence)
            current_len += sentence_len + 1  # +1 for the space

        # Flush remaining sentences
        if current_sentences:
            chunk_text_str = ' '.join(current_sentences).strip()
            if chunk_text_str:
                chunks.append({
                    "text": chunk_text_str,
                    "chunk_type": chunk_type,
                    "page_num": page_num
                })

    return chunks


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    Legacy character-based chunker — kept for backward compatibility.
    New code should prefer chunk_text_smart().
    """
    if not text or not text.strip():
        return []

    chunks = []
    text = text.strip()
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            boundary = -1
            for offset in range(overlap):
                pos = end - offset
                if pos < len(text) and text[pos].isspace():
                    boundary = pos
                    break
            if boundary != -1:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
        if start >= len(text) or chunk_size <= overlap:
            break

    return chunks


# ── PDF Table Extraction ──────────────────────────────────────────────────────

def _table_to_markdown(table: List[List[Any]]) -> str:
    """
    Convert a pdfplumber table (list of rows, each a list of cell values)
    to GitHub-flavored Markdown table format.

    Example output:
    | Name | Revenue | Growth |
    |------|---------|--------|
    | AAPL | $394B   | 8%     |
    """
    if not table:
        return ""

    # Clean cells: replace None with empty string, strip whitespace
    cleaned = []
    for row in table:
        cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
        cleaned.append(cleaned_row)

    if not cleaned:
        return ""

    # Determine column widths for alignment
    num_cols = max(len(row) for row in cleaned)

    # Pad all rows to same column count
    padded = [row + [""] * (num_cols - len(row)) for row in cleaned]

    lines = []
    # Header row (first row)
    header = padded[0]
    lines.append("| " + " | ".join(header) + " |")
    # Separator
    lines.append("| " + " | ".join(["---"] * num_cols) + " |")
    # Data rows
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def extract_pdf_content_via_gemini(filepath: str) -> List[Dict[str, Any]]:
    """
    OCR and multi-modal document parser using the Gemini File API.
    Transcribes page-by-page and returns structured content blocks.
    """
    client = get_gemini_client()
    if not client:
        print("[RAG] Gemini client not configured, cannot fall back to OCR.")
        return []

    print(f"[RAG] Uploading scanned PDF to Gemini File API: {os.path.basename(filepath)}...")
    try:
        pdf_file = client.files.upload(file=filepath)
        print(f"[RAG] Upload complete. Temporary name: {pdf_file.name}")

        # Wait for file to become active
        import time
        for _ in range(15):
            if pdf_file.state.name == "ACTIVE":
                break
            time.sleep(1)
            pdf_file = client.files.get(name=pdf_file.name)

        if pdf_file.state.name != "ACTIVE":
            print(f"[RAG] Gemini File upload stuck in state: {pdf_file.state.name}")
            return []

        print("[RAG] Transcribing scanned pages via gemini-2.5-flash...")
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                pdf_file,
                "Transcribe all text from this PDF file. Return the transcription page-by-page. For each page, start with a header like '[Page X]' followed by the transcribed text and tables. Preserve the exact content without summarizing."
            ]
        )

        # Clean up temporary file from Gemini storage
        try:
            client.files.delete(name=pdf_file.name)
            print("[RAG] Temporary file cleaned up from Gemini API storage.")
        except Exception as del_err:
            print(f"[RAG] Warning: failed to delete temp file {pdf_file.name}: {del_err}")

        transcription = response.text or ""
        blocks = []
        pages = re.split(r'\[Page\s+(\d+)\]', transcription)
        if len(pages) > 1:
            for i in range(1, len(pages), 2):
                page_num = int(pages[i])
                page_text = pages[i+1].strip()
                if page_text:
                    blocks.append({
                        "text": page_text,
                        "chunk_type": "text",
                        "page_num": page_num
                    })
        else:
            if transcription.strip():
                blocks.append({
                    "text": transcription.strip(),
                    "chunk_type": "text",
                    "page_num": 1
                })

        print(f"[RAG] Gemini OCR completed. Extracted {len(blocks)} pages.")
        return blocks

    except Exception as e:
        print(f"[RAG] Gemini OCR failed: {e}")
        return []


def extract_pdf_content(filepath: str) -> List[Dict[str, Any]]:
    """
    Extract text and tables from a PDF using standard parsers,
    falling back to Gemini OCR if it's a scanned document (image-only).
    """
    blocks = _extract_raw_pdf_content(filepath)
    total_length = sum(len(b["text"]) for b in blocks if "text" in b)
    if total_length < 100:
        print(f"[RAG] Scanned or empty PDF detected (digital text length: {total_length}). Triggering Gemini OCR...")
        ocr_blocks = extract_pdf_content_via_gemini(filepath)
        if ocr_blocks:
            return ocr_blocks
    return blocks


def _extract_raw_pdf_content(filepath: str) -> List[Dict[str, Any]]:
    """
    Extract text and tables from a PDF using pdfplumber.
    """

    try:
        import pdfplumber # type: ignore
    except ImportError:
        # Graceful fallback to pypdf if pdfplumber not installed
        print("[RAG] pdfplumber not available. Falling back to pypdf text-only extraction.")
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        blocks = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                blocks.append({
                    "text": f"[Page {i+1}]\n{page_text.strip()}",
                    "chunk_type": "text",
                    "page_num": i + 1
                })
        return blocks

    content_blocks: List[Dict[str, Any]] = []

    with pdfplumber.open(filepath) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_label = f"[Page {page_num}]"

            # ── Step 1: Extract Tables ────────────────────────────────────────
            tables = page.extract_tables()
            table_bboxes = []
            for table_obj in page.find_tables():
                table_bboxes.append(table_obj.bbox)

            for table in tables:
                if not table:
                    continue
                markdown_table = _table_to_markdown(table)
                if markdown_table.strip():
                    content_blocks.append({
                        "text": f"{page_label} [TABLE]\n{markdown_table}",
                        "chunk_type": "table",
                        "page_num": page_num
                    })

            # ── Step 2: Extract Plain Text (excluding table regions) ──────────
            # Crop out table bounding boxes and extract remaining words
            remaining_page = page
            for bbox in table_bboxes:
                try:
                    # pdfplumber uses (x0, top, x1, bottom) bbox format
                    # outside_bbox returns text regions outside the given bounding box
                    remaining_page = remaining_page.outside_bbox(bbox)
                except Exception:
                    pass  # If bbox cropping fails, keep original page text

            page_text = remaining_page.extract_text()
            if not page_text or not page_text.strip():
                continue

            # ── Step 3: Identify Headers vs Body Text ─────────────────────────
            lines = page_text.strip().split('\n')
            text_lines = []

            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue

                # Heuristic: a header is a short line that is either:
                # - ALL CAPS with at least 3 chars
                # - Ends with ':' and has fewer than 80 chars
                # - Looks like a numbered section: "1. Introduction" or "CHAPTER 2:"
                is_header = (
                    (stripped.isupper() and len(stripped) >= 3 and len(stripped) <= 120)
                    or (stripped.endswith(':') and len(stripped) <= 80)
                    or bool(re.match(r'^(\d+\.?\s+|[A-Z]{2,}\s+\d+|\bCHAPTER\b|\bSECTION\b)', stripped))
                )

                if is_header:
                    # Flush any pending body text first
                    if text_lines:
                        body_text = f"{page_label}\n" + '\n'.join(text_lines)
                        content_blocks.append({
                            "text": body_text.strip(),
                            "chunk_type": "text",
                            "page_num": page_num
                        })
                        text_lines = []
                    # Add the header as its own block
                    content_blocks.append({
                        "text": f"{page_label} [HEADER] {stripped}",
                        "chunk_type": "header",
                        "page_num": page_num
                    })
                else:
                    text_lines.append(stripped)

            # Flush remaining body text
            if text_lines:
                body_text = f"{page_label}\n" + '\n'.join(text_lines)
                content_blocks.append({
                    "text": body_text.strip(),
                    "chunk_type": "text",
                    "page_num": page_num
                })

    return content_blocks


# ── File Content Extraction ────────────────────────────────────────────────────

def extract_file_content_blocks(filepath: str, filename: str) -> List[Dict[str, Any]]:
    """
    Parse files and return a list of enriched content blocks:
        [{text, chunk_type, page_num}, ...]
    Cleans all NUL (0x00) characters from text to prevent database crashes.
    """
    raw_blocks = _extract_raw_content_blocks(filepath, filename)
    for block in raw_blocks:
        if "text" in block and isinstance(block["text"], str):
            block["text"] = block["text"].replace("\x00", "").replace("\u0000", "")
    return raw_blocks

def _extract_raw_content_blocks(filepath: str, filename: str) -> List[Dict[str, Any]]:
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found at path: {filepath}")

    ext = filename.split(".")[-1].lower()


    # ── PDF ────────────────────────────────────────────────────────────────────
    if ext == "pdf":
        return extract_pdf_content(filepath)

    # ── CSV ────────────────────────────────────────────────────────────────────
    elif ext == "csv":
        rows_text = []
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            return []

        headers = [h.strip() for h in rows[0]]

        # Build a Markdown table for a preview header block
        preview_rows = rows[:6]  # Header + up to 5 preview rows
        markdown_preview = _table_to_markdown(preview_rows)
        blocks: List[Dict[str, Any]] = []

        if markdown_preview:
            blocks.append({
                "text": f"CSV File: {filename}\nPreview (first 5 rows):\n{markdown_preview}",
                "chunk_type": "table",
                "page_num": 0
            })

        # Index each row as a searchable text block
        for i, row in enumerate(rows[1:]):
            row_vals = []
            for j, val in enumerate(row):
                header = headers[j] if j < len(headers) else f"Column {j+1}"
                row_vals.append(f"{header}={val.strip()}")
            row_text = f"Row {i+1}: {', '.join(row_vals)}"
            rows_text.append(row_text)

        # Group rows into chunks of 20 for efficient indexing
        ROWS_PER_CHUNK = 20
        for start in range(0, len(rows_text), ROWS_PER_CHUNK):
            group = rows_text[start:start + ROWS_PER_CHUNK]
            blocks.append({
                "text": '\n'.join(group),
                "chunk_type": "text",
                "page_num": 0
            })

        return blocks

    # ── JSON ───────────────────────────────────────────────────────────────────
    elif ext == "json":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            data = json.load(f)

        blocks = []
        if isinstance(data, list):
            ITEMS_PER_CHUNK = 10
            for start in range(0, len(data), ITEMS_PER_CHUNK):
                group = data[start:start + ITEMS_PER_CHUNK]
                text_items = [f"Item {start+i+1}: {json.dumps(obj)}" for i, obj in enumerate(group)]
                blocks.append({
                    "text": '\n'.join(text_items),
                    "chunk_type": "text",
                    "page_num": 0
                })
        else:
            blocks.append({
                "text": json.dumps(data, indent=2),
                "chunk_type": "text",
                "page_num": 0
            })
        return blocks

    # ── DOCX (Microsoft Word) ──────────────────────────────────────────────────
    elif ext == "docx":
        blocks = []
        try:
            import docx # type: ignore
            doc = docx.Document(filepath)
            
            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                is_header = False
                if para.style and para.style.name.startswith("Heading"):
                    is_header = True
                
                blocks.append({
                    "text": text,
                    "chunk_type": "header" if is_header else "text",
                    "page_num": 0
                })

            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_cells)
                if table_rows:
                    markdown_table = _table_to_markdown(table_rows)
                    if markdown_table:
                        blocks.append({
                            "text": markdown_table,
                            "chunk_type": "table",
                            "page_num": 0
                        })
        except Exception as e:
            print(f"[RAG] Error parsing Word document '{filename}': {e}")
            blocks.append({
                "text": f"Error parsing Word document '{filename}'. Content could not be indexed due to: {str(e)}",
                "chunk_type": "text",
                "page_num": 0
            })
        return blocks

    # ── XLSX (Microsoft Excel) ─────────────────────────────────────────────────
    elif ext == "xlsx":
        blocks = []
        try:
            import openpyxl # type: ignore
            wb = openpyxl.load_workbook(filepath, data_only=True)
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    if not any(val is not None for val in row):
                        continue
                    rows.append([str(val).strip() if val is not None else "" for val in row])
                
                if not rows:
                    continue

                headers = rows[0]
                sheet_rows_text = []
                
                # Build a Markdown table for a preview header block of this sheet
                preview_rows = rows[:6]  # Header + up to 5 preview rows
                markdown_preview = _table_to_markdown(preview_rows)
                if markdown_preview:
                    blocks.append({
                        "text": f"Excel Sheet: {sheet_name} (Preview):\n{markdown_preview}",
                        "chunk_type": "table",
                        "page_num": 0
                    })

                # Index row-by-row
                for i, row in enumerate(rows[1:]):
                    row_vals = []
                    for j, val in enumerate(row):
                        header = headers[j] if j < len(headers) and headers[j] else f"Column {j+1}"
                        row_vals.append(f"{header}={val}")
                    row_text = f"Sheet: {sheet_name}, Row {i+1}: {', '.join(row_vals)}"
                    sheet_rows_text.append(row_text)

                # Chunk rows into groups of 15 for index efficiency
                ROWS_PER_CHUNK = 15
                for start in range(0, len(sheet_rows_text), ROWS_PER_CHUNK):
                    group = sheet_rows_text[start:start + ROWS_PER_CHUNK]
                    blocks.append({
                        "text": '\n'.join(group),
                        "chunk_type": "text",
                        "page_num": 0
                    })
        except Exception as e:
            print(f"[RAG] Error parsing Excel spreadsheet '{filename}': {e}")
            blocks.append({
                "text": f"Error parsing Excel spreadsheet '{filename}'. Content could not be indexed due to: {str(e)}",
                "chunk_type": "text",
                "page_num": 0
            })
        return blocks

    # ── TXT / MD / Other ──────────────────────────────────────────────────────
    else:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
        # For Markdown, detect headers
        blocks = []
        lines = raw.split('\n')
        text_lines = []

        for line in lines:
            stripped = line.strip()
            # Detect Markdown headers (#, ##, ###)
            if re.match(r'^#{1,6}\s+\S', stripped):
                if text_lines:
                    blocks.append({
                        "text": '\n'.join(text_lines).strip(),
                        "chunk_type": "text",
                        "page_num": 0
                    })
                    text_lines = []
                blocks.append({
                    "text": stripped,
                    "chunk_type": "header",
                    "page_num": 0
                })
            else:
                text_lines.append(line)

        if text_lines:
            remaining = '\n'.join(text_lines).strip()
            if remaining:
                blocks.append({
                    "text": remaining,
                    "chunk_type": "text",
                    "page_num": 0
                })

        return blocks


def extract_file_content(filepath: str, filename: str) -> str:
    """
    Legacy flat-text extractor — kept for backward compatibility.
    Prefer extract_file_content_blocks() for new code.
    """
    blocks = extract_file_content_blocks(filepath, filename)
    return '\n\n'.join(b["text"] for b in blocks)


# ── Indexing ───────────────────────────────────────────────────────────────────

def index_file(filepath: str, filename: str, username: str = None) -> Dict[str, Any]: # type: ignore
    """
    Extract file content into enriched blocks, apply smart chunking,
    generate embeddings, and index into ChromaDB or pgvector.

    Each indexed chunk carries metadata:
        filename, chunk_index, total_chunks, username, chunk_type, page_num
    """
    if username is None:
        username = active_user_context.get()

    chunk_size_val = 800
    chunk_overlap_val = 100
    if username:
        try:
            from db import get_user_id, get_user_settings
            uid = get_user_id(username)
            if uid:
                settings = get_user_settings(uid)
                chunk_size_val = settings.get("chunk_size", 800)
                chunk_overlap_val = settings.get("chunk_overlap", 100)
        except Exception as e:
            print(f"[RAG] Error loading user settings for chunking: {e}")

    # 1. Clean up any existing index for this file
    delete_file_index(filename, username=username)

    # 2. Extract content blocks
    print(f"[RAG] Indexing file: {filename} for user: {username}...")
    content_blocks = extract_file_content_blocks(filepath, filename)
    if not content_blocks:
        return {"filename": filename, "status": "empty", "chunks": 0}

    # 3. Apply smart chunking to each block
    all_chunks: List[Dict[str, Any]] = []
    for block in content_blocks:
        block_text = block["text"]
        block_type = block.get("chunk_type", "text")
        block_page = block.get("page_num", 0)

        if block_type == "table":
            # Tables are kept as single chunks (don't split tables mid-row)
            if block_text.strip():
                all_chunks.append({
                    "text": block_text.strip(),
                    "chunk_type": "table",
                    "page_num": block_page
                })
        elif block_type == "header":
            # Headers are always their own chunk
            if block_text.strip():
                all_chunks.append({
                    "text": block_text.strip(),
                    "chunk_type": "header",
                    "page_num": block_page
                })
        else:
            # Text blocks get sentence-aware chunking
            sub_chunks = chunk_text_smart(
                block_text,
                chunk_size=chunk_size_val,
                overlap_sentences=max(1, chunk_overlap_val // 80),
                chunk_type=block_type,
                page_num=block_page
            )
            all_chunks.extend(sub_chunks)

    if not all_chunks:
        return {"filename": filename, "status": "empty", "chunks": 0}

    chunk_texts = [c["text"] for c in all_chunks]
    total = len(chunk_texts)
    print(f"[RAG] Created {total} smart chunks for '{filename}'. Generating embeddings...")

    # 4. Generate embeddings — use zero-vectors as fallback if Gemini is unavailable
    embeddings_available = True
    try:
        embeddings = get_gemini_embeddings_batch(chunk_texts)
    except Exception as emb_err:
        print(f"[RAG] Warning: Embedding generation failed for '{filename}': {emb_err}")
        print(f"[RAG] Storing chunks without embeddings — keyword search will still work.")
        embeddings_available = False
        # Use zero-vectors so the documents are stored and retrievable by keyword
        EMBEDDING_DIM = 768  # text-embedding-004 dimension
        embeddings = [[0.0] * EMBEDDING_DIM for _ in chunk_texts]

    # 5. Insert into vector store with enriched metadata
    if is_postgres_active():
        conn = get_conn()
        cur = get_cursor(conn)
        try:
            for i, (chunk, emb) in enumerate(zip(all_chunks, embeddings)):
                execute_sql(
                    cur,
                    """INSERT INTO document_embeddings
                       (filename, chunk_index, total_chunks, username, content, embedding, chunk_type, page_num)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                    (
                        filename, i, total, username,
                        chunk["text"], str(emb),
                        chunk.get("chunk_type", "text"),
                        chunk.get("page_num", 0)
                    )
                )
            conn.commit()
        except Exception as e:
            conn.rollback()
            raise e
        finally:
            conn.close()
    else:
        collection = get_chroma_collection()
        if collection is None:
            raise RuntimeError("ChromaDB collection is not initialized.")
        ids = [f"{filename}_chunk_{i}" for i in range(total)]
        metadatas = []
        for i, chunk in enumerate(all_chunks):
            meta: Dict[str, Any] = {
                "filename": filename,
                "chunk_index": i,
                "total_chunks": total,
                "chunk_type": chunk.get("chunk_type", "text"),
                "page_num": chunk.get("page_num", 0)
            }
            meta["username"] = username or "guest"
            metadatas.append(meta)

        embeddings_any: Any = embeddings
        metadatas_any: Any = metadatas
        collection.add(
            embeddings=embeddings_any,
            documents=chunk_texts,
            metadatas=metadatas_any,
            ids=ids
        )


    print(f"[RAG] Successfully indexed '{filename}' with {total} chunks ({sum(1 for c in all_chunks if c.get('chunk_type')=='table')} table chunks, {sum(1 for c in all_chunks if c.get('chunk_type')=='header')} header chunks).")
    status = "success" if embeddings_available else "keyword_only"
    return {"filename": filename, "status": status, "chunks": total}


# ── Deletion ───────────────────────────────────────────────────────────────────

def delete_file_index(filename: str, username: str = None): # type: ignore
    """
    Remove all document chunks for a specific file from ChromaDB or pgvector.
    """
    if username is None:
        username = active_user_context.get()

    if is_postgres_active():
        conn = get_conn()
        cur = get_cursor(conn)
        try:
            # Enforce strict user isolation for deletion
            execute_sql(
                cur,
                "DELETE FROM document_embeddings WHERE filename = ? AND username = ?",
                (filename, username or "guest")
            )
            conn.commit()
            print(f"[RAG] Deleted pgvector index for file: {filename} under user: {username or 'guest'}")
        except Exception as e:
            conn.rollback()
            print(f"[RAG] Error deleting pgvector index for file {filename}: {e}")
        finally:
            conn.close()
    else:
        collection = get_chroma_collection()
        if collection is None:
            return
        try:
            # Enforce strict user isolation for deletion
            where_clause = {
                "$and": [
                    {"filename": filename},
                    {"username": username or "guest"}
                ]
            }
            collection.delete(where=where_clause)
            print(f"[RAG] Deleted index for file: {filename} under user: {username or 'guest'}")
        except Exception as e:
            print(f"[RAG] Error deleting index for file {filename}: {e}")


# ── Listing ────────────────────────────────────────────────────────────────────

def get_indexed_files(username: str = None) -> List[Dict[str, Any]]: # type: ignore
    """
    Get a list of all indexed files and their chunk counts.
    """
    if username is None:
        username = active_user_context.get()

    if is_postgres_active():
        conn = get_conn()
        cur = get_cursor(conn)
        try:
            # Enforce strict user isolation for listing
            execute_sql(
                cur,
                "SELECT filename, COUNT(*) as chunks FROM document_embeddings WHERE username = ? GROUP BY filename",
                (username or "guest",)
            )
            rows = cur.fetchall()
            res = []
            for row in rows:
                try:
                    fname = row["filename"]
                    chunks = row["chunks"]
                except Exception:
                    fname = row[0]
                    chunks = row[1]
                res.append({"filename": fname, "chunks": chunks})
            return res
        except Exception as e:
            print(f"[RAG] Error listing pgvector indexed files: {e}")
            return []
        finally:
            conn.close()
    else:
        collection = get_chroma_collection()
        if collection is None:
            return []
        try:
            # Enforce strict user isolation for listing
            where_clause = {"username": username or "guest"}
            results = collection.get(where=where_clause, include=["metadatas"])
            if results is None:
                return []
            metadatas = results.get("metadatas") or []
            file_counts: Dict[str, int] = {}
            for meta_raw in metadatas:
                meta: Any = meta_raw
                if meta and isinstance(meta, dict):
                    fname = meta.get("filename")
                    if fname:
                        file_counts[fname] = file_counts.get(fname, 0) + 1
            return [{"filename": fname, "chunks": count} for fname, count in file_counts.items()]
        except Exception as e:
            print(f"[RAG] Error listing indexed files: {e}")
            return []


# ── Hybrid Search Helpers (FTS + RRF) ─────────────────────────────────────────

def _fts_search(query: str, username: str, top_k: int, cur, active_file: str = None) -> List[Dict[str, Any]]: # type: ignore
    """
    Full-Text Search (FTS) using PostgreSQL to_tsvector / plainto_tsquery.
    Returns rows ranked by text relevance, enriched with the same fields as
    the vector search so they can be merged via RRF.
    Optionally restricted to a specific file via active_file.
    """
    try:
        # Build WHERE clause dynamically
        where_parts = ["to_tsvector('english', content) @@ plainto_tsquery('english', ?)"]
        params_fts: list = []
        params_fts.append(query)  # for ts_rank
        params_fts.append(query)  # for @@ operator

        # Enforce strict user isolation (force guest scoping for unauthenticated sessions)
        where_parts.append("username = ?")
        params_fts.append(username or "guest")
        if active_file:
            where_parts.append("filename = ?")
            params_fts.append(active_file)

        where_sql = " AND ".join(where_parts)
        params_fts.append(top_k)  # for LIMIT

        execute_sql(
            cur,
            f"""
            SELECT filename, chunk_index, total_chunks, content,
                   ts_rank(to_tsvector('english', content),
                           plainto_tsquery('english', ?)) * 100 AS similarity,
                   chunk_type, page_num
            FROM document_embeddings
            WHERE {where_sql}
            ORDER BY similarity DESC
            LIMIT ?
            """,
            tuple(params_fts)
        )
        rows = cur.fetchall()
        results = []
        for row in rows:
            try:
                results.append({
                    "content": str(row["content"]),
                    "filename": str(row["filename"]),
                    "chunk_index": int(row["chunk_index"]),
                    "total_chunks": int(row["total_chunks"]) if row["total_chunks"] is not None else 0,
                    "similarity_score": round(float(row["similarity"]), 2),
                    "chunk_type": str(row["chunk_type"]) if row["chunk_type"] is not None else "text",
                    "page_num": int(row["page_num"]) if row["page_num"] is not None else 0,
                })
            except Exception:
                results.append({
                    "content": str(row[3]),
                    "filename": str(row[0]),
                    "chunk_index": int(row[1]),
                    "total_chunks": int(row[2]) if row[2] is not None else 0,
                    "similarity_score": round(float(row[4]), 2),
                    "chunk_type": str(row[5]) if row[5] is not None else "text",
                    "page_num": int(row[6]) if row[6] is not None else 0,
                })
        return results
    except Exception as e:
        print(f"[RAG] FTS search failed (non-fatal, using vector-only): {e}")
def _local_fts_search(query: str, username: str, top_k: int, active_file: str = None) -> List[Dict[str, Any]]:
    """
    Keyword search in ChromaDB by fetching metadata and documents for the active user,
    ranking using term occurrence frequency, and formatting standard output.
    """
    collection = get_chroma_collection()
    if collection is None:
        return []
    try:
        where_clause: Dict[str, Any] = {"username": username or "guest"}
        if active_file:
            where_clause = {"$and": [{"username": username or "guest"}, {"filename": active_file}]}

        all_docs = collection.get(where=where_clause, include=["documents", "metadatas"])
        if not all_docs:
            return []

        docs_list = all_docs.get("documents") or []
        metas_list = all_docs.get("metadatas") or []

        query_words = [w.lower() for w in query.split() if len(w) > 2]
        if not query_words:
            query_words = [query.lower()]

        scored = []
        for doc, meta_raw in zip(docs_list, metas_list):
            if not doc:
                continue
            doc_lower = doc.lower()
            # Rank based on query terms frequency
            score = sum(doc_lower.count(w) for w in query_words)
            if score > 0:
                meta: Any = meta_raw
                scored.append((score, doc, meta))

        scored.sort(key=lambda x: x[0], reverse=True)
        results = []
        for score, doc, meta in scored[:top_k]:
            results.append({
                "content": str(doc),
                "filename": meta.get("filename", "unknown") if hasattr(meta, "get") else "unknown",
                "chunk_index": meta.get("chunk_index", 0) if hasattr(meta, "get") else 0,
                "total_chunks": meta.get("total_chunks", 0) if hasattr(meta, "get") else 0,
                "similarity_score": round(min(score * 10.0, 100.0), 2),
                "chunk_type": meta.get("chunk_type", "text") if hasattr(meta, "get") else "text",
                "page_num": meta.get("page_num", 0) if hasattr(meta, "get") else 0,
            })
        return results
    except Exception as e:
        print(f"[RAG] Local keyword search failed: {e}")
        return []



def _reciprocal_rank_fusion(
    vector_results: List[Dict[str, Any]],
    fts_results: List[Dict[str, Any]],
    k: int = 60
) -> List[Dict[str, Any]]:
    """
    Merge two ranked lists using Reciprocal Rank Fusion (RRF).

    RRF score = sum(1 / (k + rank)) across all lists a document appears in.
    Documents that appear in BOTH lists are boosted to the top.
    k=60 is the standard constant that dampens the effect of high ranks.

    Returns a single de-duplicated list sorted by descending RRF score.
    """
    scores: Dict[str, float] = {}
    # Key = (filename, chunk_index) for de-duplication
    best: Dict[str, Dict[str, Any]] = {}

    for rank, result in enumerate(vector_results, start=1):
        key = f"{result['filename']}::{result['chunk_index']}"
        scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank)
        best[key] = result

    for rank, result in enumerate(fts_results, start=1):
        key = f"{result['filename']}::{result['chunk_index']}"
        scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank)
        if key not in best:
            best[key] = result

    merged = sorted(scores.keys(), key=lambda key: scores[key], reverse=True)
    return [best[key] for key in merged]


# ── Search ─────────────────────────────────────────────────────────────────────

def search_knowledge(query: str, top_k: int = 5) -> List[Dict[str, Any]]:
    """
    Embed the user query and retrieve the top_k most semantically similar
    document chunks from ChromaDB or pgvector.

    If active_file_context is set (document workspace mode), results are
    filtered to only include chunks from that specific file.

    Returns enriched result dicts:
        {content, filename, chunk_index, total_chunks, similarity_score, chunk_type, page_num}
    """
    username = active_user_context.get()
    active_file = active_file_context.get()  # None means all files

    if active_file:
        print(f"[RAG] Document mode: restricting search to file '{active_file}'")

    # Get query embedding — fall back to keyword search if embedding fails
    query_embedding = None
    embedding_available = True
    try:
        query_embedding = get_gemini_embeddings_batch([query])[0]
    except Exception as e:
        print(f"[RAG] Warning: Could not generate query embedding (Gemini key missing?): {e}")
        print(f"[RAG] Falling back to keyword-based search for query: '{query}'")
        embedding_available = False

    # ── Keyword fallback path (no Gemini embeddings) ──────────────────────────
    if not embedding_available:
        try:
            if is_postgres_active():
                conn = get_conn()
                cur = get_cursor(conn)
                try:
                    fts_results = _fts_search(query, username, top_k, cur, active_file)  # type: ignore
                    return fts_results
                finally:
                    conn.close()
            else:
                # ChromaDB keyword scan — fetch all docs for this user and filter
                collection = get_chroma_collection()
                if collection is None:
                    return []
                where_clause: Dict[str, Any] = {"username": username or "guest"}
                if active_file:
                    where_clause = {"$and": [{"username": username or "guest"}, {"filename": active_file}]}
                all_docs = collection.get(where=where_clause, include=["documents", "metadatas"])
                if not all_docs:
                    return []
                docs_list = all_docs.get("documents") or []
                metas_list = all_docs.get("metadatas") or []
                query_words = [w.lower() for w in query.split() if len(w) > 2]
                scored = []
                for doc, meta_raw in zip(docs_list, metas_list):
                    if not doc:
                        continue
                    score = sum(doc.lower().count(w) for w in query_words)
                    if score > 0:
                        meta: Any = meta_raw
                        scored.append((score, doc, meta))
                scored.sort(key=lambda x: x[0], reverse=True)
                results_kw = []
                for score, doc, meta in scored[:top_k]:
                    results_kw.append({
                        "content": str(doc),
                        "filename": meta.get("filename", "unknown") if hasattr(meta, "get") else "unknown",
                        "chunk_index": meta.get("chunk_index", 0) if hasattr(meta, "get") else 0,
                        "total_chunks": meta.get("total_chunks", 0) if hasattr(meta, "get") else 0,
                        "similarity_score": round(min(score * 10.0, 100.0), 2),
                        "chunk_type": meta.get("chunk_type", "text") if hasattr(meta, "get") else "text",
                        "page_num": meta.get("page_num", 0) if hasattr(meta, "get") else 0,
                    })
                return results_kw
        except Exception as kw_err:
            print(f"[RAG] Keyword fallback search failed: {kw_err}")
            return []


    if is_postgres_active():
        conn = get_conn()
        cur = get_cursor(conn)
        try:
            # ── Step 1: Semantic (vector cosine) search ──────────────────────
            # Build WHERE clause dynamically based on username and active_file
            where_clauses = []
            params_vec: list = [str(query_embedding)]

            # Enforce strict user isolation (force guest scoping for unauthenticated sessions)
            where_clauses.append("username = ?")
            params_vec.append(username or "guest")
            if active_file:
                where_clauses.append("filename = ?")
                params_vec.append(active_file)

            where_sql = ("WHERE " + " AND ".join(where_clauses)) if where_clauses else ""
            params_vec += [str(query_embedding), top_k]

            execute_sql(
                cur,
                f"""
                SELECT filename, chunk_index, total_chunks, content,
                       (1 - (embedding <=> ?::vector)) * 100 AS similarity,
                       chunk_type, page_num
                FROM document_embeddings
                {where_sql}
                ORDER BY embedding <=> ?::vector
                LIMIT ?
                """,
                tuple(params_vec)
            )
            rows = cur.fetchall()
            vector_results = []
            for row in rows:
                try:
                    vector_results.append({
                        "content": str(row["content"]),
                        "filename": str(row["filename"]),
                        "chunk_index": int(row["chunk_index"]),
                        "total_chunks": int(row["total_chunks"]) if row["total_chunks"] is not None else 0,
                        "similarity_score": round(float(row["similarity"]), 2),
                        "chunk_type": str(row["chunk_type"]) if row["chunk_type"] is not None else "text",
                        "page_num": int(row["page_num"]) if row["page_num"] is not None else 0
                    })
                except Exception:
                    vector_results.append({
                        "content": str(row[3]),
                        "filename": str(row[0]),
                        "chunk_index": int(row[1]),
                        "total_chunks": int(row[2]) if row[2] is not None else 0,
                        "similarity_score": round(float(row[4]), 2),
                        "chunk_type": str(row[5]) if row[5] is not None else "text",
                        "page_num": int(row[6]) if row[6] is not None else 0
                    })

            # ── Step 2: Full-Text Search (FTS) for exact keyword matches ─────
            fts_results = _fts_search(query, username, top_k, cur, active_file) # type: ignore

            # ── Step 3: Merge with Reciprocal Rank Fusion (RRF) ─────────────
            if fts_results:
                merged = _reciprocal_rank_fusion(vector_results, fts_results)
                print(f"[RAG] Hybrid search: {len(vector_results)} vector + "
                      f"{len(fts_results)} FTS -> {len(merged)} merged (RRF)"
                      + (f" [file={active_file}]" if active_file else ""))
                return merged[:top_k]

            # FTS found nothing — fall through to pure vector results
            return vector_results
        except Exception as e:
            print(f"[RAG] Error searching pgvector: {e}")
            return []
        finally:
            conn.close()
    else:
        collection = get_chroma_collection()
        if collection is None:
            print("[RAG] Warning: ChromaDB collection not initialized.")
            return []
        try:
            query_kwargs: Dict[str, Any] = {
                "query_embeddings": [query_embedding],
                "n_results": top_k,
                "include": ["documents", "metadatas", "distances"]
            }
            # Build ChromaDB where filter combining username and active_file
            chroma_filters = []
            # Enforce strict user isolation (force guest scoping for unauthenticated sessions)
            chroma_filters.append({"username": username or "guest"})
            if active_file:
                chroma_filters.append({"filename": active_file})

            if len(chroma_filters) == 1:
                query_kwargs["where"] = chroma_filters[0]
            elif len(chroma_filters) > 1:
                query_kwargs["where"] = {"$and": chroma_filters}

            results = collection.query(**query_kwargs)
            if results is None:
                return []

            formatted_results = []
            documents = results.get("documents")
            metadatas = results.get("metadatas")
            distances = results.get("distances")

            docs_list = documents[0] if documents and len(documents) > 0 else []
            metas_list = metadatas[0] if metadatas and len(metadatas) > 0 else []
            dists_list = distances[0] if distances and len(distances) > 0 else []

            for doc, meta_raw, dist in zip(docs_list, metas_list, dists_list):
                meta: Any = meta_raw
                if doc is not None and meta is not None and dist is not None:
                    similarity = round(max(0.0, 1.0 - (float(dist) / 2.0)) * 100, 2)
                    formatted_results.append({
                        "content": str(doc),
                        "filename": meta.get("filename", "unknown") if hasattr(meta, "get") else "unknown",
                        "chunk_index": meta.get("chunk_index", 0) if hasattr(meta, "get") else 0,
                        "total_chunks": meta.get("total_chunks", 0) if hasattr(meta, "get") else 0,
                        "similarity_score": similarity,
                        "chunk_type": meta.get("chunk_type", "text") if hasattr(meta, "get") else "text",
                        "page_num": meta.get("page_num", 0) if hasattr(meta, "get") else 0
                    })

            # Run local keyword search for hybrid matching
            fts_results = _local_fts_search(query, username, top_k, active_file)
            if fts_results:
                merged = _reciprocal_rank_fusion(formatted_results, fts_results)
                print(f"[RAG] Local hybrid search: {len(formatted_results)} vector + "
                      f"{len(fts_results)} keyword -> {len(merged)} merged (RRF)"
                      + (f" [file={active_file}]" if active_file else ""))
                return merged[:top_k]

            return formatted_results
        except Exception as e:
            print(f"[RAG] Error searching ChromaDB: {e}")
            return []
