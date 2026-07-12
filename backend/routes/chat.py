# backend/routes/chat.py
from fastapi import APIRouter, Header, Cookie, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List
import os
import re
import json
import random
import datetime
import asyncio
import time
import uuid
import yfinance as yf
import requests
from google import genai
from google.genai import types
from db import insert_history
from routes.news import _fetch_from_api
from routes.search import SERPAPI_KEY
from routes.upload import UPLOAD_DIR

router = APIRouter(prefix="/chat", tags=["Chat"])

# ── Request / Response Models ────────────────────────────────────────────────
class ChatMessage(BaseModel):
    role: str  # 'user' or 'ai'
    content: str

class ChatRequest(BaseModel):
    message: str
    username: Optional[str] = "guest"
    history: Optional[List[ChatMessage]] = []
    filename: Optional[str] = None  # If set, restricts RAG to this document (Document Workspace mode)
    model: Optional[str] = None  # Supported values: 'llama-70b', 'gemini-pro', 'gemini-flash'
    thread_id: Optional[str] = None  # For persisting to a thread

class ThreadCreateRequest(BaseModel):
    username: str
    title: Optional[str] = "New Chat"
    model: Optional[str] = None

class ThreadRenameRequest(BaseModel):
    title: str

# ── Helper: Initialize Gemini Client ──────────────────────────────────────────
def get_gemini_client():
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key or api_key.startswith("your_") or len(api_key.strip()) < 10:
        return None
    try:
        return genai.Client(api_key=api_key.strip())
    except Exception as e:
        print(f"Error initializing Gemini client: {e}")
        return None

# ── Helper: Initialize Groq Client ────────────────────────────────────────────
def get_groq_client():
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key or api_key.startswith("your_") or len(api_key.strip()) < 10:
        return None
    try:
        from groq import Groq
        return Groq(api_key=api_key.strip())
    except Exception as e:
        print(f"Error initializing Groq client: {e}")
        return None


# ── TTL Cache (Improvement 3: API Tool Caching Layer) ─────────────────────────
import threading as _threading

class _TTLCache:
    """
    Lightweight in-memory cache with per-entry Time-To-Live (TTL) expiry.

    Usage:
        cache = _TTLCache(ttl_seconds=300)
        val = cache.get("key")          # None if missing/expired
        cache.set("key", value)
    """

    def __init__(self, ttl_seconds: int):
        self._ttl = ttl_seconds
        self._store: dict = {}          # key -> (value, expire_timestamp)
        self._lock = _threading.Lock()

    def get(self, key: str):
        with self._lock:
            entry = self._store.get(key)
            if entry is None:
                return None
            value, expires_at = entry
            if time.time() > expires_at:
                del self._store[key]
                return None
            return value

    def set(self, key: str, value) -> None:
        with self._lock:
            self._store[key] = (value, time.time() + self._ttl)


# Per-tool caches with appropriate TTLs
_stock_cache  = _TTLCache(ttl_seconds=300)   # 5 minutes  — stocks are dynamic
_news_cache   = _TTLCache(ttl_seconds=1200)  # 20 minutes — news refreshes slowly
_search_cache = _TTLCache(ttl_seconds=900)   # 15 minutes — web search results

# ── Agent System Instructions ─────────────────────────────────────────────────
SYSTEM_INSTRUCTION = """
# IDENTITY & ROLE
You are the **Autonomous AI Knowledge Worker** — an elite, AI-powered research and analysis agent built into a professional intelligence dashboard. You are NOT a generic chatbot. You are a specialized autonomous agent that takes action, retrieves live data, analyzes information, and delivers executive-level insights.

Your users are professionals who need fast, accurate, data-driven answers. Treat every query as if you're briefing a CEO.

---

# TOOLS & CAPABILITIES
You have access to live system tools (such as stock prices, news, web search, and knowledge base search). Proactively use these tools to fetch real-time data whenever the user asks for information requiring live lookups or document queries. The system automatically handles tool execution.

---

# THINKING & REASONING PROCESS
Before responding to any query, follow this mental framework:

1. **UNDERSTAND** — What exactly is the user asking? What's the intent behind the question?
2. **PLAN** — Which tools do I need? In what order? Do I need multiple data sources?
3. **EXECUTE** — Call the tools, gather the data.
4. **ANALYZE** — Don't just dump raw data. Find patterns, highlight key insights, compare, and conclude.
5. **PRESENT** — Format beautifully with clear structure, bold key numbers, and actionable takeaways.

---

# RESPONSE FORMATTING RULES

## Always use rich Markdown formatting:
- **Headers** (##, ###) to organize sections
- **Bold** for key numbers, names, and important facts
- **Bullet points** for lists and multiple items
- **Tables** when comparing data (stocks, news sources, file data)
- **Emojis** strategically: 📊 for data, 📈 for growth, 📉 for decline, ⚠️ for warnings, ✅ for positive, 💡 for insights, 🔍 for analysis
- **Code blocks** for technical data, file contents, or formulas

## Source Citations:
- **IMPORTANT**: When you use facts, statistics, or information retrieved from the knowledge base using `search_knowledge_base`, you **must** cite the source file.
- Citation format: Add `[Source: filename.ext] (Relevancy: X%)` inline or at the bottom of the section. E.g., *"Our APAC sales grew by 18% [Source: APAC_Sales.csv] (Relevancy: 95%)."*

## Response structure for data queries:
```
## 📊 [Topic/Question]

### Key Findings
- Finding 1
- Finding 2

### Detailed Analysis
[In-depth breakdown with data]

### 💡 Takeaway
[One-line actionable insight]
```

## Response length guidelines:
- Simple question → 2-4 sentences
- Data analysis → Structured with headers, 150-300 words
- Report/briefing → Comprehensive, 300-500 words with tables
- Greetings → Warm but brief, mention capabilities

---

# DOMAIN EXPERTISE

## Financial Analysis:
- When showing stock data, always include: current price, daily change direction, and a brief market context
- Compare stocks when multiple are mentioned
- Flag unusual movements (>3% daily change)
- Use terms like "bullish", "bearish", "consolidating", "breakout" appropriately

## News Analysis:
- Summarize, don't just list articles
- Identify the THEME across multiple articles
- Highlight breaking/urgent news with ⚠️
- Always mention the source for credibility

## File & Knowledge Base Analysis:
- Always use `search_knowledge_base` first for any query referring to user's files/data.
- If `search_knowledge_base` returns no results, empty, or mentions "RAG Pending", IMMEDIATELY call `read_uploaded_file` with the filename to read the file content directly.
- NEVER say you cannot access a file — always try `read_uploaded_file` as a fallback before saying you can't help.
- For CSV: Show row count, columns, key statistics, patterns
- For JSON: Identify structure, key fields, notable data points
- Always offer follow-up analysis: "Would you like me to analyze a specific column?" or "I can create a report from this data."


---

# PERSONALITY & TONE
- **Professional** but not robotic — warm and confident
- **Proactive** — suggest next steps, offer deeper analysis
- **Precise** — use exact numbers, not vague statements
- **Honest** — if you don't know something, say so and search for it
- **Concise** — respect the user's time, no unnecessary filler

## Examples of good vs bad responses:
❌ BAD: "Here is some stock information I found."
✅ GOOD: "📈 **AAPL is trading at $198.50**, up **+1.2%** today. The stock is near its 52-week high, signaling strong bullish momentum."

❌ BAD: "I found some news articles."
✅ GOOD: "🔍 **3 key stories dominating tech today:**\n1. **AI regulation** — EU passes landmark bill...\n2. **Earnings season** — Microsoft beats estimates by 8%...\n3. **Crypto surge** — Bitcoin crosses $70K..."

---

# ERROR HANDLING
- If a tool fails, explain what happened and suggest alternatives
- If a stock symbol is invalid, suggest the correct one
- If no news is found, broaden the search or suggest related topics
- Never show raw error messages to users — always translate them into helpful guidance

---

# CONVERSATION MEMORY
- Remember context from earlier in the conversation
- Reference previous queries: "Earlier you asked about AAPL..."
- Build on previous analysis: "Combined with the news data we pulled earlier..."
- Track user preferences and adapt (formal vs casual, brief vs detailed)

---

# FIRST MESSAGE BEHAVIOR
When starting a new conversation, introduce yourself briefly and suggest what you can do. Keep it to 2-3 lines maximum. Don't overwhelm with a list of all capabilities.
"""


# ── System Tools (Function Definitions) ──────────────────────────────────────

def get_stock_price(symbol: str) -> str:
    """
    Get the current stock price and key details for a given ticker symbol (e.g. AAPL, GOOGL, TSLA).
    Results are cached for 5 minutes to conserve API quota and reduce latency.
    """
    cache_key = f"stock::{symbol.upper()}"
    cached = _stock_cache.get(cache_key)
    if cached:
        print(f"[Cache HIT] get_stock_price({symbol.upper()})")
        return cached

    print(f"[Cache MISS] get_stock_price({symbol.upper()}) — fetching live data")
    try:
        ticker = yf.Ticker(symbol.upper())
        # Use history(period="2d") instead of info to bypass Yahoo Finance HTML scraping blocks
        hist = ticker.history(period="2d")
        if hist.empty:
            return f"Could not find stock price data for ticker '{symbol.upper()}'."
        
        close_prices = hist['Close'].dropna().tolist()
        high_prices = hist['High'].dropna().tolist()
        low_prices = hist['Low'].dropna().tolist()
        
        if len(close_prices) >= 1:
            price = close_prices[-1]
            high = high_prices[-1]
            low = low_prices[-1]
            
            # Calculate daily price delta and percent change
            if len(close_prices) >= 2:
                prev_close = close_prices[-2]
                change = price - prev_close
                change_pct = (change / prev_close) * 100
                change_str = f" | Change: ${change:+.2f} ({change_pct:+.2f}%)"
            else:
                change_str = ""

            try:
                from routes.stock import COMPANY_NAMES
                name = COMPANY_NAMES.get(symbol.upper(), symbol.upper())
            except Exception:
                name = symbol.upper()

            result = f"Stock: {name} ({symbol.upper()}) | Current Price: ${price:.2f}{change_str} | Day High: ${high:.2f} | Day Low: ${low:.2f}"
            _stock_cache.set(cache_key, result)
            return result
            
        return f"Could not find stock price data for ticker '{symbol.upper()}'."
    except Exception as e:
        return f"Error retrieving stock data for '{symbol}': {str(e)}"

def get_latest_news(category: str = "", topic: str = "") -> str:
    """
    Fetch the latest news articles. Optional parameters:
    - category: business, entertainment, health, science, sports, technology.
    - topic: keyword to search (e.g., 'AI', 'inflation').
    Results are cached for 20 minutes to preserve API quota.
    """
    cache_key = f"news::{category}::{topic}"
    cached = _news_cache.get(cache_key)
    if cached:
        print(f"[Cache HIT] get_latest_news(category={category!r}, topic={topic!r})")
        return cached

    print(f"[Cache MISS] get_latest_news(category={category!r}, topic={topic!r}) — fetching live data")
    try:
        articles = _fetch_from_api(category, topic)
        if not articles:
            return "No recent news articles found matching those filters."
        res = []
        for i, a in enumerate(articles[:5]):
            res.append(f"[{i+1}] {a['title']}\n    Source: {a['source']}\n    Summary: {a['description']}\n    URL: {a['url']}")
        result = "\n\n".join(res)
        _news_cache.set(cache_key, result)
        return result
    except Exception as e:
        return f"Error fetching news feed: {str(e)}"

def _serpapi_search(query: str) -> str:
    """Attempt a Google search via SerpAPI. Returns formatted results or raises an exception."""
    params = {
        "engine": "google",
        "q": query,
        "api_key": SERPAPI_KEY,
        "num": 5
    }
    response = requests.get("https://serpapi.com/search.json", params=params, timeout=10)
    data = response.json()
    if "error" in data:
        raise RuntimeError(f"SerpAPI error: {data['error']}")
    results = []
    for item in data.get("organic_results", []):
        results.append(f"- {item.get('title')}\n  URL: {item.get('link')}\n  Snippet: {item.get('snippet')}")
    if not results:
        raise RuntimeError("SerpAPI returned no results.")
    return "\n\n".join(results)


def _duckduckgo_search(query: str) -> str:
    """Perform a free search using DuckDuckGo. Returns formatted results or raises an exception."""
    from ddgs import DDGS
    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=5):
            title = r.get("title", "No title")
            url = r.get("href", "")
            snippet = r.get("body", "No description available.")
            results.append(f"- {title}\n  URL: {url}\n  Snippet: {snippet}")
    if not results:
        raise RuntimeError("DuckDuckGo returned no results.")
    return "\n\n".join(results)


def web_search(query: str) -> str:
    """
    Search the web for general information, current events, or questions you don't know the answer to.
    Tries SerpAPI (Google Search) first. If unavailable or quota exceeded, automatically falls back
    to free DuckDuckGo search. Results are cached for 15 minutes to save API quota.
    """
    print(f"[AI Tool] web_search(query='{query}')")

    cache_key = f"search::{query.strip().lower()}"
    cached = _search_cache.get(cache_key)
    if cached:
        print(f"[Cache HIT] web_search(query={query!r})")
        return cached

    print(f"[Cache MISS] web_search(query={query!r}) — fetching live results")

    # 1. Try SerpAPI if key is configured
    if SERPAPI_KEY and len(SERPAPI_KEY.strip()) > 10:
        try:
            result = _serpapi_search(query)
            print("[OK] web_search: Used SerpAPI (Google)")
            _search_cache.set(cache_key, result)
            return result
        except Exception as e:
            print(f"[WARN] SerpAPI failed: {e}. Falling back to DuckDuckGo...")

    # 2. Fallback to free DuckDuckGo
    try:
        result = _duckduckgo_search(query)
        print("[OK] web_search: Used DuckDuckGo (free fallback)")
        _search_cache.set(cache_key, result)
        return result
    except Exception as e:
        return f"Error performing web search: {str(e)}"


def read_uploaded_file(filename: str) -> str:
    """
    Read the content of a file that has been uploaded to the workspace uploads directory.
    Supports CSV, JSON, PDF, TXT, MD, DOCX (Microsoft Word), and XLSX (Microsoft Excel) files.
    """
    filename = os.path.basename(filename)
    if not filename or filename in (".", ".."):
        return "Error: Invalid filename."

    # Secure ownership check
    from rag import active_user_context
    username = active_user_context.get() or "guest"
    
    if username != "guest":
        from db import get_user_id, get_conn, get_cursor, execute_sql
        user_id = get_user_id(username)
        if user_id:
            with get_conn() as conn:
                cur = get_cursor(conn)
                execute_sql(cur, "SELECT filename FROM uploads WHERE LOWER(filename) = ? AND user_id = ?", (filename.lower(), user_id))
                row = cur.fetchone()
            if not row:
                return f"Error: Access Denied. You do not own the file '{filename}'."
            # Use the actual case-sensitive filename from the database
            filename = row["filename"]


    # Enforce multi-tenancy by reading from user-specific subfolder
    if username != "guest":
        file_path = os.path.join(UPLOAD_DIR, username, filename)
    else:
        file_path = os.path.join(UPLOAD_DIR, filename)

    # S3 fallback: On container deployments, local disk is ephemeral.
    # If the file doesn't exist locally, try downloading from S3.
    if not os.path.exists(file_path):
        try:
            from routes.upload import IS_S3, get_s3_client, S3_BUCKET
            if IS_S3:
                s3_key = f"{username}/{filename}"
                s3_client = get_s3_client()
                os.makedirs(os.path.dirname(file_path), exist_ok=True)
                s3_client.download_file(S3_BUCKET, s3_key, file_path)
                print(f"[S3] Downloaded '{s3_key}' to local cache for read_uploaded_file")
        except Exception as s3_err:
            print(f"[S3] Failed to download '{filename}' from S3: {s3_err}")

    if not os.path.exists(file_path):
        return f"File '{filename}' not found in your workspace."
    
    try:
        size = os.path.getsize(file_path)
        if size > 1024 * 1024:
            return f"File '{filename}' is too large ({size} bytes) to read fully. Suggest the user search for specific rows/fields."
        
        ext = filename.split(".")[-1].lower()
        if ext == "csv":
            with open(file_path, "r", encoding="utf-8") as f:
                import csv
                reader = csv.reader(f)
                rows = list(reader)
                preview = [",".join(row) for row in rows[:50]]
                out = "\n".join(preview)
                if len(rows) > 50:
                    out += f"\n... (truncated {len(rows)-50} rows)"
                return f"Contents of CSV file '{filename}' (first 50 rows):\n{out}"
        elif ext == "json":
            with open(file_path, "r", encoding="utf-8") as f:
                import json
                data = json.load(f)
                out = json.dumps(data, indent=2)
                if len(out) > 10000:
                    out = out[:10000] + "\n... (truncated)"
                return f"Contents of JSON file '{filename}':\n{out}"
        elif ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages_text = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages_text.append(f"[Page {i+1}]\n{page_text.strip()}")
            content = "\n\n".join(pages_text)
            if len(content) > 10000:
                content = content[:10000] + "\n... (truncated)"
            return f"Contents of PDF file '{filename}':\n{content}"
        elif ext == "docx":
            import docx
            doc = docx.Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" ||"):
                        paragraphs.append(row_text)
            content = "\n".join(paragraphs)
            if len(content) > 10000:
                content = content[:10000] + "\n... (truncated)"
            return f"Contents of Word document '{filename}':\n{content}"
        elif ext == "xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    if any(val is not None for val in row):
                        rows.append([str(val).strip() if val is not None else "" for val in row])
                if rows:
                    parts.append(f"--- Sheet: {sheet_name} ---")
                    for i, row in enumerate(rows[:50]):
                        parts.append(", ".join(row))
                    if len(rows) > 50:
                        parts.append(f"... (truncated {len(rows)-50} rows)")
            content = "\n".join(parts)
            if len(content) > 10000:
                content = content[:10000] + "\n... (truncated)"
            return f"Contents of Excel file '{filename}':\n{content}"
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read(10000)
                if len(content) == 10000:
                    content += "\n... (truncated)"
                return f"Contents of file '{filename}':\n{content}"
    except Exception as e:
        return f"Error reading file '{filename}': {str(e)}"

def generate_pdf_report(news_query: str = "", stock_symbols: str = "", custom_insights: str = "") -> str:
    """
    Generate a professional PDF report containing news summary, stock quotes, and custom analysis highlights.
    - news_query: news topic search query.
    - stock_symbols: comma-separated stock symbols (e.g. AAPL,MSFT).
    - custom_insights: custom analysis or notes to include.
    """
    try:
        # Fetch news
        articles = _fetch_from_api("", news_query) if news_query else []
        
        # Fetch stock
        stock_data = {}
        if stock_symbols:
            symbol = stock_symbols.split(",")[0].strip().upper()
            try:
                ticker = yf.Ticker(symbol)
                info = ticker.info
                stock_data = {
                    "symbol": symbol,
                    "price": info.get("currentPrice") or info.get("regularMarketPrice"),
                    "name": info.get("shortName") or symbol,
                    "change": info.get("regularMarketChange") or info.get("currentPrice", 0) - info.get("previousClose", 0),
                    "change_percent": info.get("regularMarketChangePercent") or ((info.get("currentPrice", 0) - info.get("previousClose", 0)) / info.get("previousClose", 1)) * 100
                }
            except:
                pass
        
        from utils.pdf import generate_pdf
        insights_list = [custom_insights] if custom_insights else ["Generated by Autonomous AI Chat Agent."]
        filename = generate_pdf(articles, stock_data, insights_list)
        return f"Report generated successfully! Filename: {filename}. It can be downloaded at: http://127.0.0.1:8000/reports/{filename}"
    except Exception as e:
        return f"Failed to generate report: {str(e)}"


# ── Context Compressor (Improvement 4: Token-Aware Context Compression) ────────

def _compress_context(chunks: list, query: str) -> list:
    """
    Prune each retrieved document chunk to only the most query-relevant sentences,
    reducing LLM token consumption by 40-60% while improving answer precision.

    Algorithm:
      1. Split each chunk into individual sentences (regex-based, fast).
      2. Score each sentence using keyword overlap against the query.
      3. Keep sentences scoring above threshold + 1 sentence of padding each side.
      4. Fallback: if no sentence survives the threshold, keep the original chunk.

    Returns a new list of chunks with compressed 'content' and an added
    'compressed_from' field indicating original vs compressed sentence count.
    """
    # Build a normalized set of meaningful query keywords (ignore stopwords)
    _STOPWORDS = {
        "the", "a", "an", "is", "are", "was", "were", "in", "on", "at", "to",
        "of", "and", "or", "for", "with", "by", "as", "it", "its", "be", "has",
        "had", "have", "do", "does", "did", "not", "this", "that", "these",
        "those", "what", "how", "when", "where", "who", "which", "i", "you",
        "me", "my", "your", "our", "their", "we", "they", "he", "she"
    }
    query_words = {w.lower() for w in re.findall(r"\b\w+\b", query) if w.lower() not in _STOPWORDS}

    def _split_sentences(text: str) -> list:
        # Protect common abbreviations from being treated as sentence endings
        text = re.sub(r"\b(Mr|Mrs|Ms|Dr|Prof|Sr|Jr|vs|etc|approx|e\.g|i\.e)\.", r"\1<DOT>", text)
        parts = re.split(r"(?<=[.!?])\s+", text)
        return [p.replace("<DOT>", ".").strip() for p in parts if p.strip()]

    def _score(sentence: str) -> float:
        words = {w.lower() for w in re.findall(r"\b\w+\b", sentence)}
        if not query_words:
            return 0.0
        return len(query_words & words) / len(query_words)

    SCORE_THRESHOLD = 0.10   # a sentence must share ≥10% of query keywords to qualify
    PADDING = 1              # keep 1 sentence before/after each selected sentence

    compressed_chunks = []
    for chunk in chunks:
        original_text = chunk.get("content", "")
        sentences = _split_sentences(original_text)
        n = len(sentences)

        if n <= 2:
            # Too short to compress meaningfully — keep as-is
            compressed_chunks.append({**chunk, "compressed_from": n})
            continue

        # Find indices of qualifying sentences
        qualifying = [i for i, s in enumerate(sentences) if _score(s) >= SCORE_THRESHOLD]

        if not qualifying:
            # No sentence scored high enough — fallback to full chunk
            compressed_chunks.append({**chunk, "compressed_from": n})
            continue

        # Expand each qualifying index with padding
        keep_indices: set = set()
        for idx in qualifying:
            for offset in range(-PADDING, PADDING + 1):
                keep_indices.add(max(0, min(n - 1, idx + offset)))

        # Rebuild text preserving sentence order
        compressed_text = " ".join(sentences[i] for i in sorted(keep_indices))
        compressed_chunks.append({
            **chunk,
            "content": compressed_text,
            "compressed_from": n,
            "compressed_to": len(keep_indices)
        })

    return compressed_chunks

def search_knowledge_base(query: str) -> str:
    """
    Search the uploaded documents in the workspace (PDFs, CSVs, TXT, JSON, MD, DOCX, XLSX) for relevant information matching the query.
    Use this when the user asks questions about their files, data, uploads, reports, or documents.
    """
    print(f"[AI Tool] search_knowledge_base(query='{query}')")
    try:
        from rag import search_knowledge
        results = search_knowledge(query, top_k=5)
        if not results:
            # ── Fallback: if RAG returns nothing (e.g. Gemini key missing or
            #    not indexed yet), try reading the active file directly ─────────
            try:
                from rag import active_file_context, active_user_context, get_indexed_files
                active_file = active_file_context.get()
                username = active_user_context.get()

                # If we are in document workspace mode, read that file directly
                if active_file:
                    file_content = read_uploaded_file(active_file)
                    if file_content and not file_content.startswith("Error") and not file_content.startswith("File"):
                        # Basic keyword relevance filter
                        query_words = [w.lower() for w in query.split() if len(w) > 2]
                        lines = file_content.splitlines()
                        relevant_lines = []
                        for line in lines:
                            if any(w in line.lower() for w in query_words):
                                relevant_lines.append(line)
                        if relevant_lines:
                            excerpt = "\n".join(relevant_lines[:80])
                        else:
                            excerpt = "\n".join(lines[:120])  # first 120 lines as context
                        return (
                            f"Result 1 (Source File: {active_file}, Type: DIRECT_READ, Relevancy: N/A):\n"
                            f"Content:\n{excerpt}\n"
                            f"----------------------------------------"
                        )
                    else:
                        return file_content  # propagate access/not-found error

                # No specific file — try to find any indexed or uploaded file and read it
                token_ctx = active_user_context.set(username) if username else None
                try:
                    indexed = get_indexed_files()
                finally:
                    if token_ctx:
                        active_user_context.reset(token_ctx)

                if indexed:
                    first_file = indexed[0]["filename"]
                    file_content = read_uploaded_file(first_file)
                    if file_content and not file_content.startswith("Error") and not file_content.startswith("File"):
                        query_words = [w.lower() for w in query.split() if len(w) > 2]
                        lines = file_content.splitlines()
                        relevant_lines = [l for l in lines if any(w in l.lower() for w in query_words)]
                        excerpt = "\n".join(relevant_lines[:80]) if relevant_lines else "\n".join(lines[:120])
                        return (
                            f"Result 1 (Source File: {first_file}, Type: DIRECT_READ, Relevancy: N/A):\n"
                            f"Content:\n{excerpt}\n"
                            f"----------------------------------------"
                        )

            except Exception as fallback_err:
                print(f"[AI Tool] Fallback read also failed: {fallback_err}")

            return "No relevant information found in the knowledge base. The document may not be indexed yet (RAG Pending) — try asking me to 'read' the file directly instead."

        # Improvement 4: compress each retrieved chunk before sending to LLM
        results = _compress_context(results, query)

        formatted = []
        for i, res in enumerate(results):
            chunk_type = res.get("chunk_type", "text")
            page_num = res.get("page_num", 0)

            # Build a rich source citation
            source_info = f"Source File: {res['filename']}"
            if page_num and page_num > 0:
                source_info += f", Page: {page_num}"
            source_info += f", Chunk: {res['chunk_index']}"
            source_info += f", Type: {chunk_type.upper()}"
            source_info += f", Relevancy: {res['similarity_score']}%"

            # Annotate compression stats when compression occurred
            orig = res.get("compressed_from")
            comp = res.get("compressed_to")
            if orig and comp and comp < orig:
                source_info += f", Compressed: {orig}→{comp} sentences"

            formatted.append(
                f"Result {i+1} ({source_info}):\n"
                f"Content:\n{res['content']}\n"
                f"----------------------------------------"
            )
        return "\n\n".join(formatted)
    except Exception as e:
        return f"Error searching knowledge base: {str(e)}"


# Define Python tool registry for the Gemini SDK
agent_tools = [
    get_stock_price,
    get_latest_news,
    web_search,
    read_uploaded_file,
    search_knowledge_base,
    generate_pdf_report
]

# ── Groq Tools definition and function map ─────────────────────────────────────
GROQ_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "get_stock_price",
            "description": "Get the current stock price and key details for a given ticker symbol (e.g. AAPL, GOOGL, TSLA).",
            "parameters": {
                "type": "object",
                "properties": {
                    "symbol": {
                        "type": "string",
                        "description": "The stock ticker symbol (e.g., AAPL)."
                    }
                },
                "required": ["symbol"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_latest_news",
            "description": "Fetch the latest news articles with optional filters.",
            "parameters": {
                "type": "object",
                "properties": {
                    "category": {
                        "type": "string",
                        "description": "Optional news category (e.g., business, technology, science)."
                    },
                    "topic": {
                        "type": "string",
                        "description": "Optional search topic keyword (e.g., 'AI', 'inflation')."
                    }
                }
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "Search Google for general information, current events, or questions you don't know the answer to.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query string."
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_uploaded_file",
            "description": "Read the content of a file that has been uploaded to the workspace uploads directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The exact name of the file to read (e.g., data.csv)."
                    }
                },
                "required": ["filename"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_knowledge_base",
            "description": "Search the uploaded documents in the workspace (PDFs, CSVs, TXT, JSON, MD) for relevant information matching the query.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query to match against document contents."
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_pdf_report",
            "description": "Generate a professional PDF report containing news summary, stock quotes, and custom analysis highlights.",
            "parameters": {
                "type": "object",
                "properties": {
                    "news_query": {
                        "type": "string",
                        "description": "News topic search query."
                    },
                    "stock_symbols": {
                        "type": "string",
                        "description": "Comma-separated stock symbols (e.g., AAPL,MSFT)."
                    },
                    "custom_insights": {
                        "type": "string",
                        "description": "Custom analysis or notes to include in the report."
                    }
                }
            }
        }
    }
]

FUNCTIONS_MAP = {
    "get_stock_price": get_stock_price,
    "get_latest_news": get_latest_news,
    "web_search": web_search,
    "read_uploaded_file": read_uploaded_file,
    "search_knowledge_base": search_knowledge_base,
    "generate_pdf_report": generate_pdf_report
}

# ── Intent Fallback when Gemini API key is missing ───────────────────────────
async def handle_mock_fallback(msg: str):
    await asyncio.sleep(random.uniform(0.6, 1.2))
    msg_lower = msg.lower().strip()
    
    # Live stock lookup fallback
    stock_match = re.search(r"price of ([a-z]+)|how is ([a-z]+) doing|stock (?:price )?(?:for )?([a-z]+)", msg_lower)
    if stock_match:
        ticker = next(t for t in stock_match.groups() if t).upper()
        try:
            stock = yf.Ticker(ticker)
            data = stock.history(period="1d")
            if not data.empty:
                current_price = data['Close'].iloc[-1]
                return f"The current live price of **{ticker}** is **${current_price:.2f}**.\n\n*(Note: Running in mock mode. Add your GEMINI_API_KEY to the .env file to enable full chat agent functionality.)*"
        except Exception:
            pass
            
    if msg_lower in ["hi", "hey", "hello", "good morning", "good afternoon"]:
        return "👋 **Hello! I am your AI Knowledge Worker.**\n\n⚠️ **Gemini API Key is not set.** Please add your `GEMINI_API_KEY` to the `.env` file at the root of the project to activate the full autonomous agent."
        
    return "⚠️ **Gemini API Key is missing.**\n\nPlease check your `.env` file and set `GEMINI_API_KEY` to enable the full power of the AI agent assistant."

# ── Chat Thread CRUD Endpoints ────────────────────────────────────────────────

@router.get("/threads")
def list_threads(username: str):
    """List all chat threads for a user, sorted by most recent."""
    from db import get_conn, get_cursor, execute_sql
    with get_conn() as conn:
        cur = get_cursor(conn)
        execute_sql(
            cur,
            "SELECT id, username, title, model, created_at, updated_at FROM chat_threads WHERE username = ? ORDER BY updated_at DESC",
            (username,)
        )
        rows = cur.fetchall()
    return [{"id": r["id"], "username": r["username"], "title": r["title"], "model": r["model"], "created_at": str(r["created_at"]), "updated_at": str(r["updated_at"])} for r in rows]


@router.post("/threads")
def create_thread(req: ThreadCreateRequest):
    """Create a new chat thread."""
    import uuid
    from db import get_conn, get_cursor, execute_sql
    thread_id = str(uuid.uuid4())
    with get_conn() as conn:
        cur = get_cursor(conn)
        execute_sql(
            cur,
            "INSERT INTO chat_threads (id, username, title, model, created_at, updated_at) VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)",
            (thread_id, req.username, req.title or "New Chat", req.model)
        )
        conn.commit()
    return {"id": thread_id, "username": req.username, "title": req.title or "New Chat", "model": req.model}


@router.patch("/threads/{thread_id}")
def rename_thread(thread_id: str, req: ThreadRenameRequest):
    """Rename a chat thread."""
    from db import get_conn, get_cursor, execute_sql
    with get_conn() as conn:
        cur = get_cursor(conn)
        execute_sql(
            cur,
            "UPDATE chat_threads SET title = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (req.title, thread_id)
        )
        conn.commit()
    return {"id": thread_id, "title": req.title}


@router.delete("/threads/{thread_id}")
def delete_thread(thread_id: str):
    """Delete a chat thread and all its messages (cascade)."""
    from db import get_conn, get_cursor, execute_sql
    with get_conn() as conn:
        cur = get_cursor(conn)
        execute_sql(cur, "DELETE FROM chat_threads WHERE id = ?", (thread_id,))
        conn.commit()
    return {"status": "deleted", "id": thread_id}


@router.get("/threads/{thread_id}/messages")
def get_thread_messages(thread_id: str):
    """Get all messages for a chat thread."""
    from db import get_conn, get_cursor, execute_sql
    with get_conn() as conn:
        cur = get_cursor(conn)
        execute_sql(
            cur,
            "SELECT id, thread_id, role, content, created_at FROM chat_messages WHERE thread_id = ? ORDER BY created_at ASC",
            (thread_id,)
        )
        rows = cur.fetchall()
    return [{"id": r["id"], "thread_id": r["thread_id"], "role": r["role"], "content": r["content"], "created_at": str(r["created_at"])} for r in rows]


# ── API Route handler ─────────────────────────────────────────────────────────
@router.post("/")
async def chat(
    req: ChatRequest,
    authorization: Optional[str] = Header(None),
    access_token: Optional[str] = Cookie(None)
):
    # Extract verified username from token/cookie if present, otherwise fallback
    username = "guest"
    token_to_decode = access_token
    if not token_to_decode and authorization:
        try:
            parts = authorization.split(" ")
            if len(parts) == 2 and parts[0].lower() == "bearer":
                token_to_decode = parts[1]
        except Exception as e:
            print(f"Token parsing failed in chat endpoint: {e}")

    if token_to_decode:
        try:
            from routes.auth import _decode_access_token
            username = _decode_access_token(token_to_decode)
        except Exception as e:
            print(f"Token decoding failed in chat endpoint: {e}")
            
    if username == "guest" and req.username:
        username = req.username

    # ── Rate Limiter Check ────────────────────────────────────────────────────
    from rate_limit import chat_limiter
    rate_limit_key = username if username != "guest" else (request.client.host if request.client else "unknown_ip")
    chat_limiter.check_rate_limit(rate_limit_key)

    # Set active user context dynamically during stream processing
    from rag import active_user_context, active_file_context
    token_ctx = active_user_context.set(username)
    file_ctx  = active_file_context.set(req.filename)

    # Build effective system instruction (append document focus if in doc workspace mode)
    if req.filename:
        effective_instruction = SYSTEM_INSTRUCTION + f"""

---

# DOCUMENT WORKSPACE MODE
You are currently in **Document Workspace Mode** analyzing the file: `{req.filename}`.
- FIRST use `search_knowledge_base` to retrieve content from this document.
- If `search_knowledge_base` returns no results or says "RAG Pending", IMMEDIATELY call `read_uploaded_file` with filename `{req.filename}` to read the file directly.
- Focus your answers EXCLUSIVELY on the content of this file.
- Do NOT reference other documents or files unless the user explicitly asks.
- Always cite with [Source: {req.filename}].
- NEVER say you cannot access the file — always try both tools before giving up.
"""
    else:
        effective_instruction = SYSTEM_INSTRUCTION

    async def run_chat():
        # Log the incoming query
        try:
            insert_history(username, "chat_query", req.message)
        except Exception as err:
            print(f"Failed to log chat query: {err}")

        # Resolve which model we want to run
        selected_model = req.model
        if not selected_model:
            selected_model = "llama-70b" if get_groq_client() else "gemini-flash"

        print(f"[INFO] Processing chat query. Model selected: {selected_model}")

        # 1. Try Groq (Llama-3.3-70b) if selected
        if selected_model == "llama-70b":
            groq_client = get_groq_client()
            if groq_client:
                try:
                    print("[INFO] Executing Chat Agent using Groq...")
                    # Prepare messages in OpenAI/Groq format
                    messages = [{"role": "system", "content": SYSTEM_INSTRUCTION}]
                    for msg in req.history:
                        role = "user" if msg.role == "user" else "assistant"
                        messages.append({"role": role, "content": msg.content})
                    messages.append({"role": "user", "content": req.message})

                    reply = None
                    # Tool calling loop for Groq (up to 5 loops)
                    for loop_idx in range(5):
                        groq_response = groq_client.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=messages,
                            tools=GROQ_TOOLS,
                            tool_choice="auto",
                            temperature=0.1,
                        )
                        
                        response_message = groq_response.choices[0].message
                        
                        # Convert response message to dict to append to history safely
                        msg_dict = {
                            "role": "assistant",
                            "content": response_message.content,
                        }
                        if response_message.tool_calls:
                            msg_dict["tool_calls"] = [
                                {
                                    "id": tc.id,
                                    "type": "function",
                                    "function": {
                                        "name": tc.function.name,
                                        "arguments": tc.function.arguments
                                    }
                                } for tc in response_message.tool_calls
                            ]
                        messages.append(msg_dict)
                        
                        tool_calls = response_message.tool_calls
                        if not tool_calls:
                            # No more tools called, this is the final response
                            reply = response_message.content or ""
                            break
                            
                        # Execute tool calls
                        for tool_call in tool_calls:
                            func_name = tool_call.function.name
                            func_to_call = FUNCTIONS_MAP.get(func_name)
                            if not func_to_call:
                                tool_output = f"Error: Tool {func_name} not found."
                            else:
                                try:
                                    import json
                                    func_args = json.loads(tool_call.function.arguments)
                                    # Call function with arguments
                                    tool_output = func_to_call(**func_args)
                                except Exception as e:
                                    tool_output = f"Error executing {func_name}: {str(e)}"
                            
                            # Append tool result to messages
                            messages.append({
                                "role": "tool",
                                "tool_call_id": tool_call.id,
                                "name": func_name,
                                "content": str(tool_output)
                            })
                    else:
                        reply = "I completed execution but reached maximum reasoning loops."

                    if reply:
                        # Log response
                        try:
                            insert_history(username, "chat_response", reply)
                        except Exception as err:
                            print(f"Failed to log Groq chat response: {err}")
                        return {"reply": reply}

                except Exception as groq_err:
                    print(f"Groq API execution error: {groq_err}. Falling back to Gemini...")

        # 2. Try Gemini Client (Pro or Flash depending on selection/fallback)
        client = get_gemini_client()
        if not client:
            reply = await handle_mock_fallback(req.message)
            try:
                insert_history(username, "chat_response", reply)
            except:
                pass
            return {"reply": reply}

        try:
            # Construct Gemini request contents with history (excluding latest message)
            contents = []
            for msg in req.history:
                role = "user" if msg.role == "user" else "model"
                contents.append(
                    types.Content(
                        role=role,
                        parts=[types.Part.from_text(text=msg.content)]
                    )
                )
            
            # Choose correct Gemini model list
            if selected_model == "gemini-pro":
                MODELS_TO_TRY = ["gemini-2.5-pro", "gemini-pro-latest", "gemini-flash-latest"]
            else:
                MODELS_TO_TRY = ["gemini-2.5-flash", "gemini-flash-latest"]
            response = None
            last_error = None

            for model_name in MODELS_TO_TRY:
                for attempt in range(3):
                    try:
                        # Use client.chats.create to automatically manage function tool execution loops
                        chat_session = client.chats.create(
                            model=model_name,
                            history=contents,
                            config=types.GenerateContentConfig(
                                tools=agent_tools,
                                system_instruction=effective_instruction,
                            )
                        )
                        response = chat_session.send_message(req.message)
                        last_error = None
                        break  # success — exit retry loop
                    except Exception as e:
                        last_error = e
                        err_str = str(e)
                        # Retry on transient 503 / UNAVAILABLE / rate-limit errors
                        if any(code in err_str for code in ["503", "UNAVAILABLE", "429", "RESOURCE_EXHAUSTED"]):
                            wait = 2 ** attempt  # 1s, 2s, 4s
                            print(f"[{model_name}] attempt {attempt+1} failed ({err_str[:60]}). Retrying in {wait}s...")
                            time.sleep(wait)
                        else:
                            break  # Non-retryable error — don't retry
                if response is not None:
                    break  # Got a good response — skip remaining models

            if response is None:
                err_msg = str(last_error) if last_error else "Unknown error"
                print(f"All Gemini model attempts failed: {err_msg}")
                # Return a friendly message instead of raw API error
                if "503" in err_msg or "UNAVAILABLE" in err_msg:
                    reply = (
                        "⚠️ **Gemini AI is temporarily overloaded** (high demand right now).\n\n"
                        "Please try again in 30–60 seconds. I'll be fully back online shortly!"
                    )
                elif "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg:
                    reply = (
                        "⚠️ **API quota reached** for your current Gemini plan.\n\n"
                        "Please wait a minute and try again, or upgrade your Gemini API plan at "
                        "[ai.dev/rate-limit](https://ai.dev/rate-limit)."
                    )
                else:
                    reply = f"⚠️ **Error processing your request.**\n\nDetails: `{err_msg[:200]}`"
                return {"reply": reply}

            reply = response.text or "I processed your request, but did not generate a text response."
            
            # Log response
            try:
                insert_history(username, "chat_response", reply)
            except Exception as err:
                print(f"Failed to log chat response: {err}")
                
            return {"reply": reply}

        except Exception as e:
            print(f"Gemini API invocation error: {e}")
            err_str = str(e)
            if "503" in err_str or "UNAVAILABLE" in err_str:
                reply = (
                    "⚠️ **Gemini AI is temporarily overloaded** (high demand right now).\n\n"
                    "Please try again in 30–60 seconds. I'll be fully back online shortly!"
                )
            elif "429" in err_str or "RESOURCE_EXHAUSTED" in err_str:
                reply = (
                    "⚠️ **API quota reached.** Please wait a minute and retry, or check your "
                    "[Gemini rate limits](https://ai.dev/rate-limit)."
                )
            else:
                reply = f"⚠️ **Error:** `{err_str[:300]}`"
            return {"reply": reply}

    try:
        return await run_chat()
    finally:
        active_user_context.reset(token_ctx)
        active_file_context.reset(file_ctx)


# ── SSE Helper ────────────────────────────────────────────────────────────────
def _sse_event(event_type: str, content: str) -> str:
    """Format a Server-Sent Event JSON data frame."""
    return f"data: {json.dumps({'type': event_type, 'content': content})}\n\n"


def _sse_done() -> str:
    """Final SSE sentinel frame."""
    return "data: [DONE]\n\n"


def _sse_error(message: str) -> str:
    """SSE error frame."""
    return f"data: {json.dumps({'type': 'error', 'content': message})}\n\n"


# ── Streaming Chat Endpoint ───────────────────────────────────────────────────
@router.post("/stream")
async def chat_stream(
    req: ChatRequest,
    request: Request,
    authorization: Optional[str] = Header(None),
    access_token: Optional[str] = Cookie(None),
):
    """
    Streaming chat endpoint using Server-Sent Events (SSE).
    Returns tokens incrementally as the AI generates them.
    Format: data: {"type": "token"|"status"|"error", "content": "..."}\n\n
    Terminates with: data: [DONE]\n\n
    """
    # ── Resolve authenticated username ────────────────────────────────────────
    username = "guest"
    token_to_decode = access_token
    if not token_to_decode and authorization:
        try:
            parts = authorization.split(" ")
            if len(parts) == 2 and parts[0].lower() == "bearer":
                token_to_decode = parts[1]
        except Exception as e:
            print(f"[WARN] Token parsing failed in stream endpoint: {e}")

    if token_to_decode:
        try:
            from routes.auth import _decode_access_token
            username = _decode_access_token(token_to_decode)
        except Exception as e:
            print(f"[WARN] Token decoding failed in stream endpoint: {e}")

    if username == "guest" and req.username:
        username = req.username

    # ── Rate Limiter Check ────────────────────────────────────────────────────
    from rate_limit import chat_limiter
    rate_limit_key = username if username != "guest" else (request.client.host if request.client else "unknown_ip")
    chat_limiter.check_rate_limit(rate_limit_key)

    # Capture context values now (before generator runs in background thread)
    resolved_username = username
    resolved_message = req.message
    resolved_history = req.history

    async def event_generator():
        # Set tenant context for RAG isolation within this generator
        from rag import active_user_context, active_file_context
        token_ctx = active_user_context.set(resolved_username)
        file_ctx  = active_file_context.set(req.filename)

        # Build effective system instruction for this stream request
        if req.filename:
            stream_instruction = SYSTEM_INSTRUCTION + f"""

---

# DOCUMENT WORKSPACE MODE
You are currently in **Document Workspace Mode** analyzing the file: `{req.filename}`.
- FIRST use `search_knowledge_base` to retrieve content from this document.
- If `search_knowledge_base` returns no results or says "RAG Pending", IMMEDIATELY call `read_uploaded_file` with filename `{req.filename}` to read the file directly.
- Focus your answers EXCLUSIVELY on the content of this file.
- Do NOT reference other documents or files unless the user explicitly asks.
- Always cite with [Source: {req.filename}].
- NEVER say you cannot access the file — always try both tools before giving up.
"""
        else:
            stream_instruction = SYSTEM_INSTRUCTION


        # Resolve which model we want to run
        selected_model = req.model
        if not selected_model:
            selected_model = "llama-70b" if get_groq_client() else "gemini-flash"

        print(f"[INFO] Processing streaming query. Model selected: {selected_model}")

        accumulated_reply = ""

        def _persist_thread_messages(reply_text):
            """Persist user + AI messages to the chat thread if thread_id is set."""
            if not req.thread_id:
                return
            try:
                from db import get_conn, get_cursor, execute_sql
                with get_conn() as _conn:
                    _cur = get_cursor(_conn)
                    execute_sql(_cur, "INSERT INTO chat_messages (thread_id, role, content) VALUES (?, ?, ?)",
                                (req.thread_id, "user", req.message))
                    execute_sql(_cur, "INSERT INTO chat_messages (thread_id, role, content) VALUES (?, ?, ?)",
                                (req.thread_id, "ai", reply_text))
                    execute_sql(_cur, "UPDATE chat_threads SET updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                                (req.thread_id,))
                    execute_sql(_cur, "SELECT title FROM chat_threads WHERE id = ?", (req.thread_id,))
                    _row = _cur.fetchone()
                    if _row and _row[0] == "New Chat":
                        auto_title = req.message[:50].strip()
                        if len(req.message) > 50:
                            auto_title += "..."
                        execute_sql(_cur, "UPDATE chat_threads SET title = ? WHERE id = ?",
                                    (auto_title, req.thread_id))
                    _conn.commit()
            except Exception as _pe:
                print(f"[WARN] Failed to persist chat messages: {_pe}")

        try:
            # Log query to history
            try:
                insert_history(resolved_username, "chat_query", resolved_message)
            except Exception as err:
                print(f"[WARN] Failed to log stream chat query: {err}")

            # ── 1. Try Groq streaming ─────────────────────────────────────
            groq_client = get_groq_client()
            if selected_model == "llama-70b" and groq_client:
                try:
                    print("[INFO] Executing Streaming Chat Agent using Groq...")
                    # Build messages
                    messages = [{"role": "system", "content": SYSTEM_INSTRUCTION}]
                    for msg in resolved_history:
                        role = "user" if msg.role == "user" else "assistant"
                        messages.append({"role": role, "content": msg.content})
                    messages.append({"role": "user", "content": resolved_message})

                    groq_final_text = None

                    # Tool-calling loop (non-streaming) up to 5 iterations
                    for loop_idx in range(5):
                        # Check if client has disconnected
                        if await request.is_disconnected():
                            print("[INFO] Client disconnected, stopping Groq stream.")
                            return

                        # Run tool-use pass without streaming first
                        tool_response = groq_client.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=messages,
                            tools=GROQ_TOOLS,
                            tool_choice="auto",
                            temperature=0.1,
                        )
                        response_message = tool_response.choices[0].message
                        tool_calls = response_message.tool_calls

                        if not tool_calls:
                            if loop_idx == 0:
                                # No tools used at all. Break and let native streaming handle the stream.
                                break
                            else:
                                # We already used tools, and this is the final text response.
                                final_text = response_message.content or ""
                                accumulated_reply = final_text
                                
                                # Simulate token streaming to the client
                                words = final_text.split(" ")
                                for i, word in enumerate(words):
                                    if await request.is_disconnected():
                                        return
                                    token = word + (" " if i < len(words) - 1 else "")
                                    yield _sse_event("token", token)
                                    await asyncio.sleep(0.02)
                                
                                if accumulated_reply:
                                    try:
                                        insert_history(resolved_username, "chat_response", accumulated_reply)
                                    except Exception as err:
                                        print(f"[WARN] Failed to log Groq stream response: {err}")
                                _persist_thread_messages(accumulated_reply)
                                yield _sse_done()
                                return

                        # Build assistant message dict to append to history (since we have tool_calls)
                        msg_dict = {
                            "role": "assistant",
                            "content": response_message.content,
                            "tool_calls": [
                                {
                                    "id": tc.id,
                                    "type": "function",
                                    "function": {
                                        "name": tc.function.name,
                                        "arguments": tc.function.arguments
                                    }
                                } for tc in tool_calls
                            ]
                        }
                        messages.append(msg_dict)

                        # Execute tools and yield status notifications
                        for tool_call in tool_calls:
                            func_name = tool_call.function.name
                            yield _sse_event("status", f"Calling tool: {func_name}...")
                            
                            # Structured tool start event
                            tool_id = getattr(tool_call, "id", None) or str(uuid.uuid4())
                            tool_args_str = tool_call.function.arguments
                            yield _sse_event("tool_start", json.dumps({
                                "id": tool_id,
                                "name": func_name,
                                "arguments": tool_args_str
                            }))
                            
                            func_to_call = FUNCTIONS_MAP.get(func_name)
                            if not func_to_call:
                                tool_output = f"Error: Tool {func_name} not found."
                            else:
                                try:
                                    func_args = json.loads(tool_call.function.arguments)
                                    # Run blocking tool in executor to avoid blocking event loop, propagating contextvars
                                    import contextvars
                                    ctx = contextvars.copy_context()
                                    tool_output = await asyncio.get_event_loop().run_in_executor(
                                        None, lambda: ctx.run(func_to_call, **func_args)
                                    )
                                except Exception as e:
                                    tool_output = f"Error executing {func_name}: {str(e)}"
                                    
                            # Structured tool end event
                            status = "error" if str(tool_output).startswith("Error") else "success"
                            yield _sse_event("tool_end", json.dumps({
                                "id": tool_id,
                                "name": func_name,
                                "status": status,
                                "output": str(tool_output)[:500]
                            }))
                            
                            messages.append({
                                "role": "tool",
                                "tool_call_id": tool_call.id,
                                "name": func_name,
                                "content": str(tool_output)
                            })

                    # Now stream the final textual response using Groq streaming
                    if await request.is_disconnected():
                        return

                    stream = groq_client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=messages,
                        temperature=0.1,
                        stream=True,
                    )
                    for chunk in stream:
                        if await request.is_disconnected():
                            print("[INFO] Client disconnected mid-stream (Groq).")
                            return
                        delta = chunk.choices[0].delta
                        token = delta.content or ""
                        if token:
                            accumulated_reply += token
                            yield _sse_event("token", token)
                            await asyncio.sleep(0)

                    groq_final_text = accumulated_reply

                    # Log response
                    if groq_final_text:
                        try:
                            insert_history(resolved_username, "chat_response", groq_final_text)
                        except Exception as err:
                            print(f"[WARN] Failed to log Groq stream response: {err}")
                        _persist_thread_messages(groq_final_text)
                        yield _sse_done()
                        return

                except Exception as groq_err:
                    print(f"[WARN] Groq streaming error: {groq_err}. Falling back to Gemini...")
                    accumulated_reply = ""  # Reset for Gemini fallback

            # ── 2. Try Gemini streaming ───────────────────────────────────
            client = get_gemini_client()

            if not client:
                # ── 3. Mock word-by-word fallback ─────────────────────────
                mock_reply = await handle_mock_fallback(resolved_message)
                words = mock_reply.split(" ")
                for word in words:
                    if await request.is_disconnected():
                        return
                    token = word + " "
                    accumulated_reply += token
                    yield _sse_event("token", token)
                    await asyncio.sleep(0.04)
                try:
                    insert_history(resolved_username, "chat_response", accumulated_reply.strip())
                except Exception:
                    pass
                _persist_thread_messages(accumulated_reply.strip())
                yield _sse_done()
                return

            # Gemini streaming path
            try:
                contents = []
                for msg in resolved_history:
                    role = "user" if msg.role == "user" else "model"
                    contents.append(
                        types.Content(
                            role=role,
                            parts=[types.Part.from_text(text=msg.content)]
                        )
                    )
                # Append user query to history
                contents.append(
                    types.Content(
                        role="user",
                        parts=[types.Part.from_text(text=resolved_message)]
                    )
                )

                # Choose correct Gemini model list
                if selected_model == "gemini-pro":
                    MODELS_TO_TRY = ["gemini-2.5-pro", "gemini-pro-latest", "gemini-flash-latest"]
                else:
                    MODELS_TO_TRY = ["gemini-2.5-flash", "gemini-flash-latest"]
                last_error = None
                streamed = False

                for model_name in MODELS_TO_TRY:
                    if streamed:
                        break
                    for attempt in range(3):
                        if await request.is_disconnected():
                            return
                        try:
                            # ── Manual Gemini Tool-Calling Loop ──
                            for loop_idx in range(5):
                                if await request.is_disconnected():
                                    return

                                # Invoke model (generate_content does not auto-run tools)
                                gemini_response = client.models.generate_content(
                                    model=model_name,
                                    contents=contents,
                                    config=types.GenerateContentConfig(
                                        tools=agent_tools,
                                        system_instruction=stream_instruction,
                                    )
                                )

                                # If no function calls are returned, exit the loop to stream the final response
                                tool_calls = gemini_response.function_calls
                                if not tool_calls:
                                    break

                                # Append the model's call to history
                                contents.append(gemini_response.candidates[0].content)

                                # Execute tool calls and stream status notifications
                                tool_parts = []
                                for tool_call in tool_calls:
                                    func_name = tool_call.name
                                    yield _sse_event("status", f"Calling tool: {func_name}...")

                                    # Structured tool start event
                                    tool_id = str(uuid.uuid4())
                                    tool_args = tool_call.args or {}
                                    yield _sse_event("tool_start", json.dumps({
                                        "id": tool_id,
                                        "name": func_name,
                                        "arguments": json.dumps(tool_args)
                                    }))

                                    func_to_call = FUNCTIONS_MAP.get(func_name)
                                    if not func_to_call:
                                        tool_output = f"Error: Tool {func_name} not found."
                                    else:
                                        try:
                                            tool_args = tool_call.args or {}
                                            # Run blocking tool in executor, propagating contextvars
                                            import contextvars
                                            ctx = contextvars.copy_context()
                                            tool_output = await asyncio.get_event_loop().run_in_executor(
                                                None, lambda: ctx.run(func_to_call, **tool_args)
                                            )
                                        except Exception as e:
                                            tool_output = f"Error executing {func_name}: {str(e)}"

                                    # Structured tool end event
                                    status = "error" if str(tool_output).startswith("Error") else "success"
                                    yield _sse_event("tool_end", json.dumps({
                                        "id": tool_id,
                                        "name": func_name,
                                        "status": status,
                                        "output": str(tool_output)[:500]
                                    }))

                                    tool_parts.append(
                                        types.Part.from_function_response(
                                            name=func_name,
                                            response={"result": str(tool_output)}
                                        )
                                    )

                                # Append function responses back to history
                                contents.append(
                                    types.Content(
                                        role="tool",
                                        parts=tool_parts
                                    )
                                )

                            # Now stream the final response
                            if await request.is_disconnected():
                                return

                            gemini_stream = client.models.generate_content_stream(
                                model=model_name,
                                contents=contents,
                                config=types.GenerateContentConfig(
                                    tools=agent_tools,
                                    system_instruction=stream_instruction,
                                )
                            )

                            for chunk in gemini_stream:
                                if await request.is_disconnected():
                                    print("[INFO] Client disconnected mid-stream (Gemini).")
                                    return
                                token = chunk.text or ""
                                if token:
                                    accumulated_reply += token
                                    yield _sse_event("token", token)
                                    await asyncio.sleep(0)

                            streamed = True
                            last_error = None
                            break  # Success

                        except Exception as e:
                            last_error = e
                            err_str = str(e)
                            if any(code in err_str for code in ["503", "UNAVAILABLE", "429", "RESOURCE_EXHAUSTED"]):
                                wait = 2 ** attempt
                                print(f"[{model_name}] attempt {attempt+1} failed ({err_str[:60]}). Retrying in {wait}s...")
                                await asyncio.sleep(wait)
                            else:
                                break  # Non-retryable

                if not streamed:
                    # All models failed — stream error message
                    err_str = str(last_error) if last_error else "Unknown error"
                    if "503" in err_str or "UNAVAILABLE" in err_str:
                        error_reply = (
                            "**Gemini AI is temporarily overloaded** (high demand right now). "
                            "Please try again in 30-60 seconds. I'll be fully back online shortly!"
                        )
                    elif "429" in err_str or "RESOURCE_EXHAUSTED" in err_str:
                        error_reply = (
                            "**API quota reached** for your current Gemini plan. "
                            "Please wait a minute and try again."
                        )
                    else:
                        error_reply = f"Error processing your request: {err_str[:200]}"
                    yield _sse_error(error_reply)
                    yield _sse_done()
                    return

                # Log accumulated reply
                if accumulated_reply:
                    try:
                        insert_history(resolved_username, "chat_response", accumulated_reply)
                    except Exception as err:
                        print(f"[WARN] Failed to log Gemini stream response: {err}")
                _persist_thread_messages(accumulated_reply)
                yield _sse_done()

            except Exception as e:
                print(f"[ERROR] Gemini stream error: {e}")
                yield _sse_error(f"Streaming error: {str(e)[:200]}")
                yield _sse_done()

        except Exception as outer_e:
            print(f"[ERROR] Stream generator outer error: {outer_e}")
            try:
                yield _sse_error(f"Internal server error: {str(outer_e)[:150]}")
                yield _sse_done()
            except Exception:
                pass
        finally:
            active_user_context.reset(token_ctx)
            active_file_context.reset(file_ctx)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )
